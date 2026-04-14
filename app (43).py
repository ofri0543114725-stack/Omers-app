import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, copy, zipfile, re

FONT_HEBREW = "David"
FONT_ENGLISH = "Times New Roman"
FONT_SIZE_HEBREW = 13
FONT_SIZE_ENGLISH = 11
PAGE_WIDTH = Inches(8.27)
PAGE_HEIGHT = Inches(11.69)
MARGIN = Inches(1)

def is_hebrew(text):
    for ch in text:
        if '\u0590' <= ch <= '\u05FF':
            return True
    return False

def has_digits(text):
    return any(ch.isdigit() for ch in text)

def split_mixed_run(run):
    """פצל run שמכיל מעורב אנגלית+ספרות לרצף של runs"""
    text = run.text
    if not text:
        return
    has_latin = any('a' <= ch.lower() <= 'z' for ch in text)
    has_digit = any(ch.isdigit() for ch in text)
    if not (has_latin and has_digit and not is_hebrew(text)):
        return  # לא מעורב, אין צורך לפצל

    # בנה קבוצות של תווים לפי סוג
    import re
    groups = re.findall(r'[a-zA-Z]+|[0-9]+|[^a-zA-Z0-9]+', text)
    if len(groups) <= 1:
        return

    # קבל את ה-parent element
    parent = run._element.getparent()
    idx = list(parent).index(run._element)

    import copy
    # צור run חדש לכל קבוצה
    new_runs = []
    for group in groups:
        new_run_elem = copy.deepcopy(run._element)
        # עדכן טקסט
        t = new_run_elem.find(qn('w:t'))
        if t is None:
            from docx.oxml import OxmlElement
            t = OxmlElement('w:t')
            new_run_elem.append(t)
        t.text = group
        if group.startswith(' ') or group.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_runs.append((new_run_elem, group))

    # הסר את ה-run המקורי והכנס את החדשים
    run._element.text = ''
    for i, (new_run_elem, group) in enumerate(new_runs):
        parent.insert(idx + i, new_run_elem)
    parent.remove(run._element)

fix_run_font = None  # replaced by _fix_font_only

def _fix_font_only(run):
    """שנה רק פונט וגודל — מספרים תמיד David 13"""
    text = run.text
    if not text:
        return
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # נקה את כל attributes של rFonts ותתחיל מחדש
    for attr in list(rFonts.attrib.keys()):
        del rFonts.attrib[attr]
    # עברית או ספרות בלבד — David 13
    # אנגלית (עם או בלי ספרות) — Times New Roman 11
    has_latin = any(('a' <= ch.lower() <= 'z') for ch in text)
    if is_hebrew(text) or (has_digits(text) and not has_latin):
        font, size = FONT_HEBREW, FONT_SIZE_HEBREW
    else:
        font, size = FONT_ENGLISH, FONT_SIZE_ENGLISH
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    for tag in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn('w:val'), str(size * 2))

def is_caption_text(text):
    s = text.strip()
    return s.startswith(("טבלה ", "שרטוט ", "תרשים "))

def convert_anchors_to_inline(doc):
    """המר תמונות צפות ל-inline כדי למנוע חפיפה"""
    import copy
    from lxml import etree

    WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    anchor_tag = f'{{{WP}}}anchor'
    inline_tag = f'{{{WP}}}inline'

    for drawing in doc.element.body.findall('.//' + qn('w:drawing')):
        anchor = drawing.find(f'{{{WP}}}anchor')
        if anchor is None:
            continue

        # צור inline חדש
        inline = OxmlElement('wp:inline')
        inline.set('distT', '0')
        inline.set('distB', '0')
        inline.set('distL', '0')
        inline.set('distR', '0')

        # העתק extent (גודל)
        extent = anchor.find(f'{{{WP}}}extent')
        if extent is not None:
            inline.append(copy.deepcopy(extent))

        # הוסף effectExtent ריק
        eff = OxmlElement('wp:effectExtent')
        eff.set('l', '0')
        eff.set('t', '0')
        eff.set('r', '0')
        eff.set('b', '0')
        inline.append(eff)

        # העתק docPr
        docPr = anchor.find(f'{{{WP}}}docPr')
        if docPr is not None:
            inline.append(copy.deepcopy(docPr))

        # העתק cNvGraphicFramePr
        cNv = anchor.find(f'{{{WP}}}cNvGraphicFramePr')
        if cNv is not None:
            inline.append(copy.deepcopy(cNv))

        # העתק graphic
        A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        graphic = anchor.find(f'{{{A}}}graphic')
        if graphic is not None:
            inline.append(copy.deepcopy(graphic))

        # החלף anchor ב-inline
        drawing.remove(anchor)
        drawing.append(inline)

def separate_caption_and_drawing(doc):
    """פצל פסקאות שמכילות כותרת + תמונה: כותרת מעל, תמונה מתחת"""
    import copy
    body = doc.element.body

    for elem in list(body):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag != 'p':
            continue

        from docx.text.paragraph import Paragraph
        try:
            para = Paragraph(elem, doc)
        except:
            continue

        text = para.text.strip()
        has_drawing = elem.find('.//' + qn('w:drawing')) is not None

        if not has_drawing or not text:
            continue

        # פסקה עם גם תמונה וגם טקסט — בדוק אם הטקסט הוא כותרת
        if not any(text.startswith(w) for w in ('שרטוט', 'טבלה', 'איור', 'תרשים')):
            continue

        # צור פסקת כותרת חדשה
        caption_elem = OxmlElement('w:p')
        # העתק pPr מהפסקה המקורית
        old_pPr = elem.find(qn('w:pPr'))
        if old_pPr is not None:
            caption_elem.append(copy.deepcopy(old_pPr))

        # הוסף את ה-runs של הטקסט לפסקת הכותרת
        for run in elem.findall(qn('w:r')):
            if run.find('.//' + qn('w:drawing')) is None:
                caption_elem.append(copy.deepcopy(run))

        # צור פסקת תמונה חדשה
        img_elem = OxmlElement('w:p')
        # הוסף pPr עם יישור מרכז
        img_pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        img_pPr.append(jc)
        img_elem.append(img_pPr)

        # הוסף רק את ה-runs עם תמונות
        for run in elem.findall(qn('w:r')):
            if run.find('.//' + qn('w:drawing')) is not None:
                img_elem.append(copy.deepcopy(run))

        # הוסף שורה ריקה בין כותרת לתמונה (רווח שורה וחצי)
        spacer = OxmlElement('w:p')
        spacer_pPr = OxmlElement('w:pPr')
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:line'), '360')
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')
        spacer_pPr.append(sp)
        spacer.append(spacer_pPr)

        # החלף את הפסקה המקורית בסדר: כותרת → רווח → תמונה
        elem.addprevious(caption_elem)
        elem.addprevious(spacer)
        # שמור רק את התמונה בפסקה המקורית
        for child in list(elem):
            tag_c = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_c == 'r' and child.find('.//' + qn('w:drawing')) is None:
                elem.remove(child)
            elif tag_c == 'pPr':
                # עדכן יישור מרכז לפסקת התמונה
                jc_existing = child.find(qn('w:jc'))
                if jc_existing is None:
                    jc_new = OxmlElement('w:jc')
                    jc_new.set(qn('w:val'), 'center')
                    child.append(jc_new)
                else:
                    jc_existing.set(qn('w:val'), 'center')

def add_semicolons_to_lists(doc):
    """הוסף ; בסוף כל שורת רשימה חוץ מהאחרונה, ו. באחרונה"""
    body = doc.element.body
    paras = list(body.iter(qn('w:p')))

    # מצא קבוצות של פסקאות עם אותו numId
    i = 0
    while i < len(paras):
        elem = paras[i]
        numPr = elem.find('.//' + qn('w:numPr'))
        if numPr is None:
            i += 1
            continue

        numId_el = numPr.find(qn('w:numId'))
        if numId_el is None:
            i += 1
            continue
        nid = numId_el.get(qn('w:val'))

        # מצא את כל הפסקאות ברצף עם אותו numId
        group = [elem]
        j = i + 1
        while j < len(paras):
            next_elem = paras[j]
            next_numPr = next_elem.find('.//' + qn('w:numPr'))
            if next_numPr is None:
                # בדוק אם פסקה ריקה — אפשר לדלג
                from docx.text.paragraph import Paragraph
                try:
                    next_para = Paragraph(next_elem, doc)
                    if not next_para.text.strip():
                        j += 1
                        continue
                except:
                    pass
                break
            next_numId_el = next_numPr.find(qn('w:numId'))
            if next_numId_el is None or next_numId_el.get(qn('w:val')) != nid:
                break
            group.append(next_elem)
            j += 1

        if len(group) < 2:
            i = j
            continue

        # הוסף ; לכולם חוץ מהאחרון, . לאחרון
        for k, p_elem in enumerate(group):
            suffix = '.' if k == len(group) - 1 else ';'
            # מצא את ה-run האחרון עם טקסט
            runs = p_elem.findall(qn('w:r'))
            last_text_run = None
            for run in reversed(runs):
                t = run.find(qn('w:t'))
                if t is not None and t.text and t.text.strip():
                    last_text_run = run
                    break
            if last_text_run is not None:
                t = last_text_run.find(qn('w:t'))
                # הסר סימנים קיימים בסוף
                text = t.text.rstrip()
                while text and text[-1] in '.;:,':
                    text = text[:-1].rstrip()
                t.text = text + suffix
                # שמור xml:space
                if ' ' in t.text:
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        i = j

def process_document(input_bytes):
    # עבוד ישירות על המסמך המקורי — לא על מסמך חדש
    import copy
    doc = Document(io.BytesIO(input_bytes))

    # הוסף ; ו. לרשימות
    add_semicolons_to_lists(doc)

    # עבור על כל הפסקאות ותקן פונט + רווח
    for para in doc.paragraphs:
        for run in list(para.runs):
            split_mixed_run(run)
        for run in para.runs:
            _fix_font_only(run)
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        # כותרות
        if is_caption_text(para.text.strip()):
            if pPr.find(qn('w:keepNext')) is None:
                kn = OxmlElement('w:keepNext')
                pPr.insert(0, kn)
            for run in para.runs:
                if run.text.strip():
                    run.bold = True
                    run.underline = True

    # תקן פסקאות ריקות — David 11
    for para in doc.paragraphs:
        if not para.text.strip() and para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is None:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                pPr.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            for attr in list(rFonts.attrib.keys()):
                del rFonts.attrib[attr]
            rFonts.set(qn('w:ascii'), FONT_HEBREW)
            rFonts.set(qn('w:hAnsi'), FONT_HEBREW)
            rFonts.set(qn('w:cs'), FONT_HEBREW)
            for tag_name in ['w:sz', 'w:szCs']:
                el = rPr.find(qn(tag_name))
                if el is None:
                    el = OxmlElement(tag_name)
                    rPr.append(el)
                el.set(qn('w:val'), str(11 * 2))

    # תקן pPr rPr לפונט David (משפיע על מספרי רשימה)
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                pPr.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            for attr in list(rFonts.attrib.keys()):
                del rFonts.attrib[attr]
            rFonts.set(qn('w:ascii'), FONT_HEBREW)
            rFonts.set(qn('w:hAnsi'), FONT_HEBREW)
            rFonts.set(qn('w:cs'), FONT_HEBREW)
            for tag_name in ['w:sz', 'w:szCs']:
                el = rPr.find(qn(tag_name))
                if el is None:
                    el = OxmlElement(tag_name)
                    rPr.append(el)
                el.set(qn('w:val'), str(FONT_SIZE_HEBREW * 2))

    # נקה שורות ריקות כפולות
    body = doc.element.body
    prev_empty = False
    to_remove = []
    for elem in list(body):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            try:
                from docx.text.paragraph import Paragraph
                para = Paragraph(elem, doc)
                text = para.text.strip()
                has_drawing = elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None
                is_empty = not text and not has_drawing
                if is_empty and prev_empty:
                    to_remove.append(elem)
                prev_empty = is_empty
            except:
                prev_empty = False
        else:
            prev_empty = False
    for elem in to_remove:
        body.remove(elem)

    # תאי טבלאות — פונט + יישור מרכז
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        split_mixed_run(run)
                    for run in para.runs:
                        _fix_font_only(run)
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        para._element.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'center')

    # עדכן פונט לכל ה-runs כולל תיבות טקסט
    for elem in doc.element.body.iter():
        if elem.tag == qn('w:r'):
            rPr = elem.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    cs = rFonts.get(qn('w:cs'))
                    ascii_f = rFonts.get(qn('w:ascii'))
                    if cs and not ascii_f:
                        rFonts.set(qn('w:ascii'), cs)
                        rFonts.set(qn('w:hAnsi'), cs)
                    for attr in list(rFonts.attrib.keys()):
                        if 'Theme' in attr or 'theme' in attr:
                            del rFonts.attrib[attr]

    # המר תמונות צפות ל-inline
    convert_anchors_to_inline(doc)

    # הפרד כותרת מתמונה
    separate_caption_and_drawing(doc)

    # יישור ימין
    fix_styles_rtl_doc(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()

def fix_styles_rtl_doc(doc):
    """תקן יישור RTL ב-docDefaults ובכל סגנונות הפסקה"""
    styles_elem = doc.element.find(qn('w:styles'))
    if styles_elem is None:
        return

    # תקן docDefaults
    docDefaults = styles_elem.find(qn('w:docDefaults'))
    if docDefaults is not None:
        pPrDefault = docDefaults.find(qn('w:pPrDefault'))
        if pPrDefault is None:
            pPrDefault = OxmlElement('w:pPrDefault')
            docDefaults.append(pPrDefault)
        pPr_def = pPrDefault.find(qn('w:pPr'))
        if pPr_def is None:
            pPr_def = OxmlElement('w:pPr')
            pPrDefault.append(pPr_def)
        for tag in [qn('w:bidi'), qn('w:jc')]:
            el = pPr_def.find(tag)
            if el is not None:
                pPr_def.remove(el)
        bidi = OxmlElement('w:bidi')
        pPr_def.insert(0, bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr_def.insert(1, jc)

    # תקן כל סגנון פסקה
    for style in styles_elem.findall(qn('w:style')):
        if style.get(qn('w:type')) != 'paragraph':
            continue
        pPr = style.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            style.append(pPr)
        for tag in [qn('w:bidi'), qn('w:jc')]:
            el = pPr.find(tag)
            if el is not None:
                pPr.remove(el)
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)


def copy_media(old_bytes, new_bytes):
    with zipfile.ZipFile(io.BytesIO(old_bytes)) as old_zip:
        old_rels = old_zip.read('word/_rels/document.xml.rels').decode('utf-8')
        img_rels = re.findall(r'<Relationship[^>]+image[^>]+/>', old_rels)
        media_files = {f: old_zip.read(f) for f in old_zip.namelist() if f.startswith('word/media/')}
        old_ct = old_zip.read('[Content_Types].xml').decode('utf-8')

    if not img_rels and not media_files:
        return new_bytes

    with zipfile.ZipFile(io.BytesIO(new_bytes)) as tmp:
        new_rels_text = tmp.read('word/_rels/document.xml.rels').decode('utf-8')
    existing_ids = set(re.findall(r'Id="(rId\d+)"', new_rels_text))

    id_map = {}
    counter = 100
    new_img_rels = []
    for rel in img_rels:
        while f'rId{counter}' in existing_ids:
            counter += 1
        old_id = re.search(r'Id="(rId\d+)"', rel).group(1)
        new_rel = re.sub(r'Id="rId\d+"', f'Id="rId{counter}"', rel)
        new_img_rels.append(new_rel)
        existing_ids.add(f'rId{counter}')
        id_map[old_id] = f'rId{counter}'
        counter += 1

    new_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(new_bytes)) as zin:
        with zipfile.ZipFile(new_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            existing = set(zin.namelist())
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml' and id_map:
                    text = data.decode('utf-8')
                    for oid, nid in id_map.items():
                        text = text.replace(f'r:embed="{oid}"', f'r:embed="{nid}"')
                    data = text.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels' and new_img_rels:
                    insert = '\n  '.join(new_img_rels)
                    data = data.decode('utf-8').replace('</Relationships>', f'  {insert}\n</Relationships>').encode('utf-8')
                elif item.filename == '[Content_Types].xml' and media_files:
                    ct = data.decode('utf-8')
                    for t in re.findall(r'<Default[^>]+image[^>]+/>', old_ct):
                        ext = re.search(r'Extension="([^"]+)"', t)
                        if ext and ext.group(1) not in ct:
                            ct = ct.replace('</Types>', f'  {t}\n</Types>')
                    data = ct.encode('utf-8')
                zout.writestr(item, data)
            for fname, fdata in media_files.items():
                if fname not in existing:
                    zout.writestr(fname, fdata)
    new_buf.seek(0)
    return new_buf.read()

def remove_custom_xml(docx_bytes):
    """הסר customXml שגורם לאזהרת 'גורם לא אמין' ב-Word"""
    new_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
        with zipfile.ZipFile(new_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename.startswith('customXml'):
                    continue
                data = zin.read(item.filename)
                if item.filename in ('_rels/.rels', 'word/_rels/document.xml.rels'):
                    text = data.decode('utf-8')
                    text = re.sub(r'<Relationship[^>]+customXml[^>]+/>', '', text)
                    data = text.encode('utf-8')
                elif item.filename == '[Content_Types].xml':
                    text = data.decode('utf-8')
                    text = re.sub(r'<Override[^>]+customXml[^>]+/>', '', text)
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    new_buf.seek(0)
    return new_buf.read()

def restore_missing_parts(old_bytes, new_bytes):
    """העתק קבצים חיוניים שחסרים בפלט מהמקור + תקן relationships"""
    new_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(old_bytes)) as old_zip:
        with zipfile.ZipFile(io.BytesIO(new_bytes)) as new_zip:
            existing = set(new_zip.namelist())
            # קבצים חיוניים שצריך להעתיק אם חסרים
            critical = [f for f in old_zip.namelist()
                       if f not in existing
                       and not f.startswith('customXml')
                       and (f.endswith('.xml') or f.endswith('.rels') or f.endswith('.jpeg') or f.endswith('.png'))]
            
            # בנה מיפוי של relationships חסרים מה-rels המקורי
            old_doc_rels = old_zip.read('word/_rels/document.xml.rels').decode('utf-8')
            new_doc_rels = new_zip.read('word/_rels/document.xml.rels').decode('utf-8')
            # מצא relationships שיש במקור אבל לא בחדש
            old_rels_items = re.findall(r'<Relationship[^>]+/>', old_doc_rels)
            new_rels_items = re.findall(r'<Relationship[^>]+/>', new_doc_rels)
            new_rels_ids = set(re.findall(r'Id="([^"]+)"', new_doc_rels))
            missing_rels = []
            for rel in old_rels_items:
                rel_id = re.search(r'Id="([^"]+)"', rel)
                if rel_id and rel_id.group(1) not in new_rels_ids:
                    # אל תוסיף customXml
                    if 'customXml' not in rel:
                        missing_rels.append(rel)

            with zipfile.ZipFile(new_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in new_zip.infolist():
                    data = new_zip.read(item.filename)
                    # הוסף relationships חסרים ל-document.xml.rels
                    if item.filename == 'word/_rels/document.xml.rels' and missing_rels:
                        text = data.decode('utf-8')
                        insert = '\n  '.join(missing_rels)
                        text = text.replace('</Relationships>', f'  {insert}\n</Relationships>')
                        data = text.encode('utf-8')
                    zout.writestr(item, data)
                for fname in critical:
                    zout.writestr(fname, old_zip.read(fname))
    new_buf.seek(0)
    return new_buf.read()

# ============================================================
st.set_page_config(page_title="עיצוב מסמכים", layout="centered")
st.title("🖊️ עיצוב מסמך אוטומטי")
st.markdown("העלה קובץ Word — המסמך יצא מעוצב לפי תקן החברה.")

uploaded = st.file_uploader("בחר קובץ Word", type=["docx"])

if uploaded:
    st.success(f"קובץ נטען: {uploaded.name}")
    if st.button("עצב מסמך"):
        with st.spinner("מעצב..."):
            try:
                result = process_document(uploaded.read())
                st.download_button(
                    label="📥 הורד מסמך מעוצב",
                    data=result,
                    file_name=f"מעוצב_{uploaded.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("המסמך מוכן!")
            except Exception as e:
                st.error(f"שגיאה: {e}")
                raise e
